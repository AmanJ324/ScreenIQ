import re
import io
from typing import List, Set, Dict, Any
import pypdf
import docx

# A comprehensive list of standard technical and soft skills to detect in text
COMMON_SKILLS = {
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "php", "go", "rust", 
    "swift", "kotlin", "scala", "r", "sql", "html", "css", "bash", "shell", "perl", "dart",
    # Frameworks & Libraries
    "fastapi", "flask", "django", "express", "node.js", "node", "react", "angular", "vue", 
    "next.js", "nextjs", "spring boot", "laravel", "rails", "jquery", "bootstrap", "tailwind",
    "redux", "graphql", "nest.js", "nestjs", "pytorch", "tensorflow", "keras", "scikit-learn", 
    "sklearn", "pandas", "numpy", "scipy", "nltk", "spacy", "opencv", "huggingface", "transformers",
    # Cloud & DevOps
    "docker", "kubernetes", "aws", "gcp", "azure", "jenkins", "git", "github", "gitlab", 
    "ci/cd", "terraform", "ansible", "linux", "unix", "nginx", "apache", "prometheus", "grafana",
    # Databases & Big Data
    "postgresql", "mysql", "mongodb", "redis", "sqlite", "oracle", "cassandra", "dynamodb", 
    "elasticsearch", "firebase", "spark", "hadoop", "hive", "kafka", "pyspark", "snowflake",
    # Concepts & Methodologies
    "agile", "scrum", "rest api", "microservices", "system design", "oop", "tdd", "ci / cd",
    "machine learning", "deep learning", "computer vision", "nlp", "natural language processing",
    "data science", "data analysis", "devops", "cloud computing", "statistics", "mathematics",
    # Soft & Management Skills
    "communication", "leadership", "project management", "problem solving", "teamwork", 
    "collaboration", "management", "critical thinking", "negotiation", "presentation"
}

EDUCATION_LEVELS = {
    "phd": 4, "doctorate": 4, "ph.d.": 4, "d.phil.": 4,
    "master": 3, "m.s.": 3, "m.sc.": 3, "m.tech": 3, "mba": 3, "m.e.": 3, "m.phil.": 3, "ms": 3,
    "bachelor": 2, "b.s.": 2, "b.sc.": 2, "b.tech": 2, "b.e.": 2, "ba": 2, "bs": 2, "bba": 2,
    "associate": 1, "a.s.": 1, "a.a.": 1,
    "high school": 0, "diploma": 0, "ged": 0
}

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF file bytes."""
    text = ""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error parsing PDF: {e}")
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX file bytes."""
    text = ""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
    return text

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Determines file type and extracts text from raw bytes."""
    filename_lc = filename.lower()
    if filename_lc.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename_lc.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    else:
        # Fallback to plain text decoding
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Error decoding text: {e}")
            return ""

def clean_text(text: str) -> str:
    """Cleans text for vectorization and similarity comparison."""
    # Convert to lowercase
    text = text.lower()
    # Replace newlines and carriage returns with spaces
    text = re.sub(r'[\r\n]+', ' ', text)
    # Remove non-alphanumeric characters but keep spaces, dots, and hyphens (for skills like node.js, c++)
    # Let's keep alphanumeric, plus, sharp/hash, dot, hyphen, space
    text = re.sub(r'[^a-zA-Z0-9\s\.\+#\-]', ' ', text)
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_skills(text: str, custom_skills: List[str] = None) -> List[str]:
    """Extracts skills from text by checking matches against a skills list."""
    text_lc = text.lower()
    found_skills = set()
    
    # Check common skills
    for skill in COMMON_SKILLS:
        # Use word boundary or specific pattern for skills that have symbols like C++ or .NET
        # Escape skill for safe regex compilation
        escaped_skill = re.escape(skill)
        # Handle special skills containing symbols like +, #, or .
        if '+' in skill or '#' in skill or '.' in skill:
            pattern = rf'(?:^|\s|/){escaped_skill}(?:$|\s|/|,|\.)'
        else:
            pattern = rf'\b{escaped_skill}\b'
            
        if re.search(pattern, text_lc):
            found_skills.add(skill)
            
    # Check custom skills if provided
    if custom_skills:
        for skill in custom_skills:
            skill_lc = skill.lower()
            escaped_skill = re.escape(skill_lc)
            if '+' in skill_lc or '#' in skill_lc or '.' in skill_lc:
                pattern = rf'(?:^|\s|/){escaped_skill}(?:$|\s|/|,|\.)'
            else:
                pattern = rf'\b{escaped_skill}\b'
            if re.search(pattern, text_lc):
                found_skills.add(skill_lc)
                
    return sorted(list(found_skills))

def extract_experience_years(text: str) -> float:
    """
    Heuristically extracts total years of experience from text using:
    1. Mentions of "X years of experience"
    2. Sum of year ranges like (2018 - 2022)
    """
    text_lc = text.lower()
    
    # 1. Look for years of experience phrases
    # Matches: "5 years of...", "3+ years", "10 yrs", "4.5 years"
    exp_phrases = re.findall(r'\b(\d+(?:\.\d+)?)\+?\s*(?:year|yr)s?\b', text_lc)
    max_mentioned = 0.0
    if exp_phrases:
        try:
            max_mentioned = max(float(val) for val in exp_phrases)
        except ValueError:
            pass
            
    # 2. Look for year-to-year or year-to-present ranges
    # Matches: "2018-2022", "2015 to 2021", "2020 - Present", "2021 - current"
    # Current year is 2026
    current_year = 2026
    
    # Find all 4-digit numbers that could be years (e.g. 1990 - 2026)
    range_pattern = r'\b(19\d{2}|20\d{2})\s*(?:\-|to|until)\s*(19\d{2}|20\d{2}|present|current|now)\b'
    ranges = re.findall(range_pattern, text_lc)
    
    total_range_years = 0.0
    for start, end in ranges:
        try:
            start_yr = int(start)
            if end in ['present', 'current', 'now']:
                end_yr = current_year
            else:
                end_yr = int(end)
                
            duration = end_yr - start_yr
            if 0 < duration <= 40: # filter out unrealistic ranges (e.g., typos)
                total_range_years += duration
        except ValueError:
            continue
            
    # Combine heuristics: use the max of either mentioned years or sum of ranges (capped logically)
    estimated_experience = max(max_mentioned, total_range_years)
    # Cap experience at 40 years to prevent outliers from parsing errors
    return min(estimated_experience, 40.0)

def extract_education(text: str) -> Dict[str, Any]:
    """Identifies the highest education level found in the text."""
    text_lc = text.lower()
    highest_level = "high school"
    highest_rank = 0
    
    # Check level matches
    for term, rank in EDUCATION_LEVELS.items():
        # Escape term
        escaped_term = re.escape(term)
        if term in ['ms', 'bs', 'ba']:
            # Require word boundary for short degree abbreviations
            pattern = rf'\b{escaped_term}\b'
        else:
            pattern = rf'{escaped_term}'
            
        if re.search(pattern, text_lc):
            if rank > highest_rank:
                highest_rank = rank
                # Standardize level string
                if rank == 4:
                    highest_level = "PhD"
                elif rank == 3:
                    highest_level = "Master"
                elif rank == 2:
                    highest_level = "Bachelor"
                elif rank == 1:
                    highest_level = "Associate"
                    
    return {
        "level": highest_level,
        "rank": highest_rank
    }

def parse_resume(file_bytes: bytes, filename: str, custom_skills: List[str] = None) -> Dict[str, Any]:
    """Parses a resume and extracts text, cleaned text, skills, experience, and education."""
    raw_text = extract_text_from_bytes(file_bytes, filename)
    cleaned = clean_text(raw_text)
    
    skills = extract_skills(raw_text, custom_skills)
    experience = extract_experience_years(raw_text)
    education = extract_education(raw_text)
    
    return {
        "filename": filename,
        "raw_text": raw_text,
        "cleaned_text": cleaned,
        "skills": skills,
        "experience": experience,
        "education_level": education["level"],
        "education_rank": education["rank"]
    }
